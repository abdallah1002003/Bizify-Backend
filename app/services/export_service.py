import json
import os
import uuid
from datetime import datetime
from typing import Any, Dict, List, Optional

from docx import Document
from fpdf import FPDF
from sqlalchemy.orm import Session

from app.core.celery_app import celery_app
from app.core.database import SessionLocal
from app.models.export_job import ExportJob, ExportStatus
from app.repositories.export_repo import export_repo
from app.repositories.idea_repo import idea_repo
from app.repositories.profile_repo import profile_repo
from app.repositories.skill_repo import skill_repo


def _sanitize_for_fpdf(text: str) -> str:
    """Replace non-latin1 characters to avoid FPDF font errors."""
    if not text:
        return ""

    text_str = str(text)
    return "".join(character if ord(character) < 256 else "?" for character in text_str)


def _generate_docx(data: Dict[str, Any], file_path: str) -> None:
    """Generate a DOCX export file."""
    document = Document()
    document.add_heading("Bizify Data Export", 0)

    if "profile" in data:
        document.add_heading("Profile Info", level=1)
        profile = data["profile"]
        document.add_paragraph(f"Bio: {profile.get('bio', '')}")
        document.add_paragraph(
            f"Interests: {json.dumps(profile.get('interests', []), ensure_ascii=False)}"
        )
        document.add_paragraph(
            f"Background: {json.dumps(profile.get('background', {}), ensure_ascii=False)}"
        )

    if "skills" in data:
        document.add_heading("Skills", level=1)
        for skill in data["skills"]:
            document.add_paragraph(f"- {skill['name']} (Level: {skill['level']})")

    if "ideas" in data:
        document.add_heading("Ideas", level=1)
        for idea in data["ideas"]:
            document.add_heading(idea.get("title", "Untitled"), level=2)
            document.add_paragraph(f"Status: {idea.get('status', '')}")
            document.add_paragraph(str(idea.get("description", "")))

    document.save(file_path)


def _generate_pdf(data: Dict[str, Any], file_path: str) -> None:
    """Generate a PDF export file."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", size=12)
    pdf.set_font("helvetica", "B", 15)
    pdf.cell(
        0,
        10,
        _sanitize_for_fpdf("Bizify Data Export"),
        border=False,
        align="C",
        new_x="LMARGIN",
        new_y="NEXT",
    )
    pdf.ln(10)

    if "profile" in data:
        pdf.set_font("helvetica", "B", 14)
        pdf.cell(0, 10, "Profile Info", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("helvetica", size=11)
        profile = data["profile"]
        bio = _sanitize_for_fpdf(f"Bio: {profile.get('bio', '')}")
        interests = _sanitize_for_fpdf(f"Interests: {json.dumps(profile.get('interests', []))}")
        background = _sanitize_for_fpdf(f"Background: {json.dumps(profile.get('background', {}))}")
        pdf.multi_cell(0, 8, f"{bio}\n{interests}\n{background}")
        pdf.ln(5)

    if "skills" in data:
        pdf.set_font("helvetica", "B", 14)
        pdf.cell(0, 10, "Skills", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("helvetica", size=11)
        for skill in data["skills"]:
            item = _sanitize_for_fpdf(f"- {skill['name']} (Level: {skill['level']})")
            pdf.cell(0, 8, item, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(5)

    if "ideas" in data:
        pdf.set_font("helvetica", "B", 14)
        pdf.cell(0, 10, "Ideas", new_x="LMARGIN", new_y="NEXT")
        for idea in data["ideas"]:
            pdf.set_font("helvetica", "B", 12)
            title = _sanitize_for_fpdf(idea.get("title", "Untitled"))
            pdf.cell(0, 8, title, new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("helvetica", "I", 10)
            status_text = _sanitize_for_fpdf(f"Status: {idea.get('status', '')}")
            pdf.cell(0, 6, status_text, new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("helvetica", size=11)
            description = _sanitize_for_fpdf(str(idea.get("description", "")))
            pdf.multi_cell(0, 6, description)
            pdf.ln(4)

    pdf.output(file_path)


class ExportService:
    """User data export workflows."""

    @staticmethod
    def request_export(
        db: Session,
        user_id: uuid.UUID,
        scope: List[str],
        file_format: str,
    ) -> ExportJob:
        """Create an export job and queue the background task."""
        job = export_repo.create(
            db,
            obj_in={
                "user_id": user_id,
                "scope": scope,
                "format": file_format,
                "status": ExportStatus.PENDING,
            },
        )
        task = process_export_task.delay(str(job.id))
        return export_repo.update(db, db_obj=job, obj_in={"task_id": task.id})

    @staticmethod
    def cancel_export(db: Session, job_id: uuid.UUID, user_id: uuid.UUID) -> bool:
        """Cancel a pending or in-progress export job."""
        job = export_repo.get_for_user(db, job_id, user_id)
        if job and job.status in [ExportStatus.PENDING, ExportStatus.PROCESSING]:
            celery_app.control.revoke(job.task_id, terminate=True)
            export_repo.update_status(db, job, ExportStatus.CANCELLED)
            return True

        return False

    @staticmethod
    def get_job_for_user(
        db: Session,
        job_id: uuid.UUID,
        user_id: uuid.UUID,
    ) -> Optional[ExportJob]:
        """Fetch an export job if it belongs to the user."""
        return export_repo.get_for_user(db, job_id, user_id)


@celery_app.task(name="app.services.export_service.process_export_task")
def process_export_task(job_id: str) -> None:
    """Collect user data and write the export artifact."""
    db = SessionLocal()
    job = None

    try:
        job = export_repo.get_by_id(db, uuid.UUID(job_id))
        if not job:
            return

        export_repo.update_status(db, job, ExportStatus.PROCESSING)
        data_to_export: Dict[str, Any] = {}
        user_id = job.user_id

        if "profile" in job.scope:
            profile = profile_repo.get_by_user_id(db, user_id)
            if profile:
                data_to_export["profile"] = {
                    "bio": profile.bio,
                    "interests": profile.interests_json,
                    "background": profile.background_json,
                }

        if "skills" in job.scope:
            skills = skill_repo.get_by_user(db, user_id)
            data_to_export["skills"] = [
                {"name": skill.skill_name, "level": skill.declared_level}
                for skill in skills
            ]

        if "ideas" in job.scope:
            ideas = idea_repo.get_by_owner(db, user_id)
            data_to_export["ideas"] = [
                {
                    "title": idea.title,
                    "description": idea.description,
                    "status": idea.status,
                }
                for idea in ideas
            ]

        storage_dir = "storage/exports"
        os.makedirs(storage_dir, exist_ok=True)

        format_type = getattr(job, "format", "pdf").lower()
        file_extension = "docx" if format_type == "word" else format_type
        file_name = f"export_{job.id}.{file_extension}"
        file_path = os.path.join(storage_dir, file_name)

        if format_type == "word":
            _generate_docx(data_to_export, file_path)
        elif format_type == "pdf":
            _generate_pdf(data_to_export, file_path)
        else:
            with open(file_path, "w", encoding="utf-8") as file_handle:
                json.dump(data_to_export, file_handle, ensure_ascii=False, indent=4)

        export_repo.update(
            db,
            db_obj=job,
            obj_in={
                "status": ExportStatus.COMPLETED,
                "storage_path": file_path,
                "completed_at": datetime.utcnow(),
            },
        )
    except Exception as exc:
        if job:
            export_repo.update(
                db,
                db_obj=job,
                obj_in={
                    "status": ExportStatus.FAILED,
                    "error_details": {"error": str(exc)},
                },
            )
    finally:
        db.close()
